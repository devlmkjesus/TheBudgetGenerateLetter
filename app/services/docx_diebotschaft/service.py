import base64
import io
import json
import os
import re
import time
import urllib.parse
import urllib.request
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.shared import Inches, Pt

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _color_from_supabase(value: str) -> Optional[WD_COLOR_INDEX]:
    if not value:
        return None
    v = value.strip().lower()
    if v == "yellow":
        return WD_COLOR_INDEX.YELLOW
    if v == "red":
        return WD_COLOR_INDEX.RED
    return None


_HIGHLIGHT_CACHE: Dict[str, Any] = {"by_client": {}}
_HIGHLIGHT_TTL_SECONDS = 300


def _fetch_spellcheck_words_for_client(client: str) -> List[Tuple[str, WD_COLOR_INDEX]]:
    supabase_url = (
        os.getenv("SUPABASE_URL")
        or os.getenv("VITE_SUPABASE_URL")
        or ""
    ).strip().rstrip("/")
    supabase_key = (
        os.getenv("SUPABASE_ANON_KEY")
        or os.getenv("VITE_SUPABASE_ANON_KEY")
        or ""
    ).strip()
    if not supabase_url or not supabase_key:
        return []

    now = time.time()
    by_client = _HIGHLIGHT_CACHE.get("by_client") or {}
    entry = by_client.get(client) or {}
    cached_ts = float(entry.get("ts") or 0.0)
    cached_items = entry.get("items") or []
    if now - cached_ts < _HIGHLIGHT_TTL_SECONDS:
        return list(cached_items)

    def _request_rows() -> List[Dict[str, Any]]:
        query = urllib.parse.urlencode({"select": "Word,HighlightColor,Client"})
        url = f"{supabase_url}/rest/v1/Spellcheck_The_Budget?{query}"

        req = urllib.request.Request(
            url,
            headers={
                "apikey": supabase_key,
                "Authorization": f"Bearer {supabase_key}",
                "Accept": "application/json",
            },
            method="GET",
        )

        try:
            with urllib.request.urlopen(req, timeout=10) as resp:
                raw = resp.read().decode("utf-8")
        except Exception:
            return []

        try:
            rows = json.loads(raw)
        except Exception:
            return []

        return rows if isinstance(rows, list) else []

    rows = _request_rows()
    items: List[Tuple[str, WD_COLOR_INDEX]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        row_client = (row.get("Client") or row.get("client") or "").strip()
        if client.lower() not in row_client.lower():
            continue
        word = (row.get("Word") or "").strip()
        color = _color_from_supabase(row.get("HighlightColor") or "")
        if word and color:
            items.append((word, color))

    if "by_client" not in _HIGHLIGHT_CACHE or not isinstance(_HIGHLIGHT_CACHE.get("by_client"), dict):
        _HIGHLIGHT_CACHE["by_client"] = {}
    _HIGHLIGHT_CACHE["by_client"][client] = {"ts": now, "items": items}
    return items


def _find_highlight_spans(
    text: str,
    words: List[Tuple[str, WD_COLOR_INDEX]],
) -> List[Tuple[int, int, WD_COLOR_INDEX]]:
    if not text or not words:
        return []

    spans: List[Tuple[int, int, WD_COLOR_INDEX]] = []
    for word, color in words:
        escaped = re.escape(word)
        if re.fullmatch(r"[A-Za-z0-9']+", word):
            pattern = re.compile(rf"\b{escaped}\b", re.IGNORECASE)
        else:
            pattern = re.compile(escaped, re.IGNORECASE)

        for m in pattern.finditer(text):
            spans.append((m.start(), m.end(), color))

    spans.sort(key=lambda s: (s[0], s[1]))
    return spans


def _set_font(run, *, name: str, size_pt: int, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def _add_formatted_runs(
    paragraph,
    text: str,
    highlight_spans: List[Tuple[int, int, WD_COLOR_INDEX]],
) -> None:
    if text is None:
        text = ""

    boundaries = {0, len(text)}
    for s, e, _ in highlight_spans:
        boundaries.add(s)
        boundaries.add(e)

    points = sorted(boundaries)

    def _highlight_at(i: int) -> Optional[WD_COLOR_INDEX]:
        for s, e, color in highlight_spans:
            if s <= i < e:
                return color
        return None

    for a, b in zip(points, points[1:]):
        if a == b:
            continue
        seg = text[a:b]
        r = paragraph.add_run(seg)
        _set_font(r, name="Times New Roman", size_pt=10, bold=False, italic=False)

        color = _highlight_at(a)
        if color is not None:
            r.font.highlight_color = color


def generate_diebotschaft_docx_bytes(
    *,
    batchNumber: Optional[str],
    state: Optional[str],
    churchDistrict: Optional[str],
    author: Optional[str],
    date: Optional[str],
    body: Any,
) -> bytes:
    safe_batch = batchNumber or ""
    safe_state = state or ""
    safe_district = churchDistrict or ""
    safe_author = author or ""
    safe_date = date or ""

    # Normalize body to string
    if isinstance(body, str):
        body_text = body
    elif isinstance(body, list):
        body_text = "\n".join(str(item) for item in body)
    else:
        body_text = str(body)

    doc = Document()

    # Section setup: portrait, Letter size, margins
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Two columns with 0.5" spacing
    from docx.oxml.ns import qn
    cols = section._sectPr.xpath("./w:cols")
    if not cols:
        cols = section._sectPr.add_new(qn("w:cols"))
    else:
        cols = cols[0]
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), str(int(Inches(0.5).twips)))

    # Add header with variables
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(0)
    header_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if safe_batch:
        r = header_para.add_run(safe_batch)
        _set_font(r, name="Arial", size_pt=9, bold=True)

    if safe_state:
        if safe_batch:
            header_para.add_run(" ")
        r = header_para.add_run(safe_state)
        _set_font(r, name="Times New Roman", size_pt=10, bold=True)

    if safe_district:
        if safe_batch or safe_state:
            header_para.add_run(" ")
        r = header_para.add_run(safe_district)
        _set_font(r, name="Times New Roman", size_pt=10, bold=True)

    # Author and Date in header (next paragraph)
    if safe_author or safe_date:
        author_date_para = header.add_paragraph()
        author_date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_date_para.paragraph_format.space_before = Pt(0)
        author_date_para.paragraph_format.space_after = Pt(0)
        author_date_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        if safe_author:
            r = author_date_para.add_run(safe_author)
            _set_font(r, name="Times New Roman", size_pt=10, bold=False, italic=True)

        if safe_date:
            if safe_author:
                author_date_para.add_run(" ")
            r = author_date_para.add_run(safe_date)
            _set_font(r, name="Times New Roman", size_pt=10, bold=False, italic=False)

    # Body paragraphs with two-column layout
    spellcheck_words = _fetch_spellcheck_words_for_client("DieBotschaft")
    lines = str(body_text).splitlines() or [""]

    for line in lines:
        p = doc.add_paragraph()
        p.style = "No Spacing"
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.right_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0.13)

        spans = _find_highlight_spans(line, spellcheck_words)
        _add_formatted_runs(p, line, spans)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_base64(data: bytes) -> str:
    return base64.b64encode(data).decode("utf-8")

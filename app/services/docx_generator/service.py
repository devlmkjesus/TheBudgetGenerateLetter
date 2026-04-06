import base64
import io
import json
import os
import re
import time
import urllib.parse
import urllib.request
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, Pt

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def parse_openai_json_content(raw_content: str) -> Any:
    if not isinstance(raw_content, str):
        raise ValueError("rawContent must be a string")

    cleaned = raw_content

    if "```" in cleaned:
        cleaned = cleaned.replace("```json", "").replace("```", "").strip()

    cleaned = cleaned.lstrip()
    if cleaned.lower().startswith("json"):
        cleaned = cleaned[4:].lstrip()

    try:
        return json.loads(cleaned)
    except Exception as e:
        raise ValueError(f"Failed to parse JSON: {str(e)}")


def normalize_body_to_string(body: Any) -> str:
    if body is None:
        return ""
    if isinstance(body, str):
        return body
    if isinstance(body, list):
        parts = []
        for item in body:
            if isinstance(item, str):
                parts.append(item)
            else:
                parts.append(json.dumps(item, ensure_ascii=False))
        return "\n\n".join(parts)
    if isinstance(body, dict):
        try:
            return json.dumps(body, ensure_ascii=False, indent=2)
        except Exception:
            return str(body)
    return str(body)


def _set_times_new_roman_12pt(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def _color_from_supabase(value: str) -> Optional[WD_COLOR_INDEX]:
    if not value:
        return None
    v = value.strip().lower()
    if v == "yellow":
        return WD_COLOR_INDEX.YELLOW
    if v == "red":
        return WD_COLOR_INDEX.RED
    return None


_HIGHLIGHT_CACHE: Dict[str, Any] = {"by_client": {}}
_HIGHLIGHT_TTL_SECONDS = 300


def _fetch_spellcheck_words_for_client(client: str) -> list[tuple[str, WD_COLOR_INDEX]]:
    supabase_url = (
        os.getenv("SUPABASE_URL")
        or os.getenv("VITE_SUPABASE_URL")
        or ""
    ).strip().rstrip("/")
    supabase_key = (
        os.getenv("SUPABASE_ANON_KEY")
        or os.getenv("VITE_SUPABASE_ANON_KEY")
        or ""
    ).strip()
    if not supabase_url or not supabase_key:
        return []

    now = time.time()
    by_client = _HIGHLIGHT_CACHE.get("by_client") or {}
    entry = by_client.get(client) or {}
    cached_ts = float(entry.get("ts") or 0.0)
    cached_items = entry.get("items") or []
    if now - cached_ts < _HIGHLIGHT_TTL_SECONDS:
        return list(cached_items)

    def _request_rows(filter_key: str) -> List[Dict[str, Any]]:
        query = urllib.parse.urlencode(
            {
                "select": "Word,HighlightColor,Client",
                filter_key: f"ilike.*{client}*",
            }
        )
        url = f"{supabase_url}/rest/v1/Spellcheck_The_Budget?{query}"

        req = urllib.request.Request(
            url,
            headers={
                "apikey": supabase_key,
                "Authorization": f"Bearer {supabase_key}",
                "Accept": "application/json",
            },
            method="GET",
        )

        try:
            with urllib.request.urlopen(req, timeout=10) as resp:
                raw = resp.read().decode("utf-8")
        except Exception:
            return []

        try:
            rows = json.loads(raw)
        except Exception:
            return []

        return rows if isinstance(rows, list) else []

    rows = _request_rows("client")
    items: list[tuple[str, WD_COLOR_INDEX]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        word = (row.get("Word") or "").strip()
        color = _color_from_supabase(row.get("HighlightColor") or "")
        if word and color:
            items.append((word, color))

    if not items:
        rows = _request_rows("Client")
        for row in rows:
            if not isinstance(row, dict):
                continue
            word = (row.get("Word") or "").strip()
            color = _color_from_supabase(row.get("HighlightColor") or "")
            if word and color:
                items.append((word, color))

    if "by_client" not in _HIGHLIGHT_CACHE or not isinstance(_HIGHLIGHT_CACHE.get("by_client"), dict):
        _HIGHLIGHT_CACHE["by_client"] = {}
    _HIGHLIGHT_CACHE["by_client"][client] = {"ts": now, "items": items}
    return items


def _find_highlight_spans(
    text: str,
    words: List[Tuple[str, WD_COLOR_INDEX]],
) -> List[Tuple[int, int, WD_COLOR_INDEX]]:
    if not text or not words:
        return []

    spans: List[Tuple[int, int, WD_COLOR_INDEX]] = []
    for word, color in words:
        escaped = re.escape(word)
        if re.fullmatch(r"[A-Za-z0-9']+", word):
            pattern = re.compile(rf"\b{escaped}\b", re.IGNORECASE)
        else:
            pattern = re.compile(escaped, re.IGNORECASE)

        for m in pattern.finditer(text):
            spans.append((m.start(), m.end(), color))

    spans.sort(key=lambda s: (s[0], s[1]))
    return spans


_DIARY_VALIDATION_PATTERN = re.compile(r"\b(?:the\s+diary|diary)\b", re.IGNORECASE)
_DIARY_MARKUP_PATTERN = re.compile(r"(\*{1,2}|_{1,2})(\s*(?:the\s+diary|diary)\s*)\1", re.IGNORECASE)


def strip_diary_markup(text: str) -> str:
    if not text:
        return text

    def _repl(match: re.Match[str]) -> str:
        return match.group(2)

    return _DIARY_MARKUP_PATTERN.sub(_repl, text)


def normalize_diary_casing(text: str) -> str:
    if not text:
        return text

    def _repl(match: re.Match[str]) -> str:
        lowered = match.group(0).lower()
        if lowered == "the diary":
            return "The Diary"
        return "Diary"

    return _DIARY_VALIDATION_PATTERN.sub(_repl, text)


def add_runs_with_diary_italics(paragraph, text: str) -> None:
    if not text:
        r = paragraph.add_run("")
        _set_times_new_roman_12pt(r)
        return

    idx = 0
    for match in _DIARY_VALIDATION_PATTERN.finditer(text):
        start, end = match.span()
        if start > idx:
            r = paragraph.add_run(text[idx:start])
            _set_times_new_roman_12pt(r)

        phrase = text[start:end]
        r_phrase = paragraph.add_run(phrase)
        r_phrase.italic = True
        _set_times_new_roman_12pt(r_phrase)
        idx = end

    if idx < len(text):
        r = paragraph.add_run(text[idx:])
        _set_times_new_roman_12pt(r)


def add_runs_with_diary_italics_and_highlight(
    paragraph,
    text: str,
    highlight_spans: List[Tuple[int, int, WD_COLOR_INDEX]],
) -> None:
    if text is None:
        text = ""

    diary_spans = [m.span() for m in _DIARY_VALIDATION_PATTERN.finditer(text)]

    boundaries = {0, len(text)}
    for s, e in diary_spans:
        boundaries.add(s)
        boundaries.add(e)
    for s, e, _ in highlight_spans:
        boundaries.add(s)
        boundaries.add(e)

    points = sorted(boundaries)

    def _in_diary(i: int) -> bool:
        for s, e in diary_spans:
            if s <= i < e:
                return True
        return False

    def _highlight_at(i: int) -> Optional[WD_COLOR_INDEX]:
        for s, e, color in highlight_spans:
            if s <= i < e:
                return color
        return None

    for a, b in zip(points, points[1:]):
        if a == b:
            continue
        seg = text[a:b]
        r = paragraph.add_run(seg)
        _set_times_new_roman_12pt(r)

        if _in_diary(a):
            r.italic = True

        color = _highlight_at(a)
        if color is not None:
            r.font.highlight_color = color


def generate_docx_bytes(*, city: Optional[str], author_name: Optional[str], date: Optional[str], body: Any) -> bytes:
    safe_city = city or "City"
    safe_author = author_name or "Name"
    safe_date = date or "No Date"

    title = f"{safe_city} – {safe_author}"
    body_text = normalize_body_to_string(body)

    doc = Document()

    # Title paragraph
    p_title = doc.add_paragraph()
    p_title.style = "No Spacing"
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_title.paragraph_format.line_spacing = Pt(12)
    r_title = p_title.add_run(title)
    r_title.bold = True
    _set_times_new_roman_12pt(r_title)

    # Body paragraphs (split by line breaks)
    full_text = f"{safe_date} – {body_text}" if body_text else f"{safe_date} – "
    full_text = strip_diary_markup(full_text)
    full_text = normalize_diary_casing(full_text)
    lines = str(full_text).splitlines() or [""]

    spellcheck_words = _fetch_spellcheck_words_for_client("Diary")

    for line in lines:
        p = doc.add_paragraph()
        p.style = "No Spacing"
        # python-docx justification support varies; this is the closest setting.
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.right_indent = Inches(0)
        p.paragraph_format.first_line_indent = Inches(0.13)
        spans = _find_highlight_spans(line, spellcheck_words)
        add_runs_with_diary_italics_and_highlight(p, line, spans)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_base64(data: bytes) -> str:
    return base64.b64encode(data).decode("utf-8")

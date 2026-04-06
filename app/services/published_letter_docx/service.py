import base64
import io
import os
import re
import time
import urllib.parse
import urllib.request
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches, Pt


DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


_DIARY_PATTERN = re.compile(r"\b(?:the\s+diary|diary)\b", re.IGNORECASE)
_DIARY_MARKUP_PATTERN = re.compile(r"(\*{1,2}|_{1,2})(\s*(?:the\s+diary|diary)\s*)\1", re.IGNORECASE)
_BUDGET_PATTERN = re.compile(r"\b(?:the\s+budget|budget)\b", re.IGNORECASE)
_BUDGET_MARKUP_PATTERN = re.compile(r"(\*{1,2}|_{1,2})(\s*(?:the\s+budget|budget)\s*)\1", re.IGNORECASE)


def to_base64(data: bytes) -> str:
    return base64.b64encode(data).decode("utf-8")


def strip_diary_markup(text: str) -> str:
    if not text:
        return text

    def _repl(match: re.Match[str]) -> str:
        return match.group(2)

    return _DIARY_MARKUP_PATTERN.sub(_repl, text)


def normalize_diary_casing(text: str) -> str:
    if not text:
        return text

    def _repl(match: re.Match[str]) -> str:
        lowered = match.group(0).lower()
        if lowered == "the diary":
            return "The Diary"
        return "Diary"

    return _DIARY_PATTERN.sub(_repl, text)


def strip_budget_markup(text: str) -> str:
    if not text:
        return text

    def _repl(match: re.Match[str]) -> str:
        return match.group(2)

    return _BUDGET_MARKUP_PATTERN.sub(_repl, text)


def normalize_budget_casing(text: str) -> str:
    if not text:
        return text

    def _repl(match: re.Match[str]) -> str:
        lowered = match.group(0).lower()
        if lowered == "the budget":
            return "The Budget"
        return "Budget"

    return _BUDGET_PATTERN.sub(_repl, text)


def _set_font(run, *, name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold


def _normalize_location_line(letter_heading: str) -> str:
    raw = (letter_heading or "").strip()
    if not raw:
        return ""

    s = " ".join(raw.replace("\t", " ").split())

    if "," in s:
        return s.upper()

    parts = s.split(" ")
    if len(parts) >= 2:
        city = " ".join(parts[:-1])
        state = parts[-1]
        return f"{city}, {state}".upper()

    return s.upper()


def _split_paragraphs(body: Any) -> List[str]:
    if body is None:
        return []
    if isinstance(body, list):
        return [str(p) for p in body]

    text = str(body).replace("\r\n", "\n").replace("\r", "\n")
    if not text.strip():
        return []

    # Treat blank lines as paragraph separators; otherwise preserve single newlines as separate paragraphs.
    if re.search(r"\n\s*\n", text):
        blocks = re.split(r"\n\s*\n", text.strip())
        out: List[str] = []
        for b in blocks:
            out.extend([line for line in b.split("\n") if line.strip()])
        return out

    return [line for line in text.split("\n") if line.strip()]


@dataclass(frozen=True)
class HighlightSpan:
    start: int
    end: int
    color: WD_COLOR_INDEX


def _color_from_supabase(value: str) -> Optional[WD_COLOR_INDEX]:
    if not value:
        return None
    v = value.strip().lower()
    if v == "yellow":
        return WD_COLOR_INDEX.YELLOW
    if v == "red":
        return WD_COLOR_INDEX.RED
    return None


_HIGHLIGHT_CACHE: Dict[str, Any] = {"by_client": {}}
_HIGHLIGHT_TTL_SECONDS = 300


def _fetch_spellcheck_words_for_client(client: str) -> List[Tuple[str, WD_COLOR_INDEX]]:
    supabase_url = (
        os.getenv("SUPABASE_URL")
        or os.getenv("VITE_SUPABASE_URL")
        or ""
    ).strip().rstrip("/")
    supabase_key = (
        os.getenv("SUPABASE_ANON_KEY")
        or os.getenv("VITE_SUPABASE_ANON_KEY")
        or ""
    ).strip()
    if not supabase_url or not supabase_key:
        return []

    now = time.time()
    by_client = _HIGHLIGHT_CACHE.get("by_client") or {}
    entry = by_client.get(client) or {}
    cached_ts = float(entry.get("ts") or 0.0)
    cached_items = entry.get("items") or []
    if now - cached_ts < _HIGHLIGHT_TTL_SECONDS:
        return list(cached_items)

    def _request_rows() -> List[Dict[str, Any]]:
        query = urllib.parse.urlencode({"select": "Word,HighlightColor,Client"})
        url = f"{supabase_url}/rest/v1/Spellcheck_The_Budget?{query}"

        req = urllib.request.Request(
            url,
            headers={
                "apikey": supabase_key,
                "Authorization": f"Bearer {supabase_key}",
                "Accept": "application/json",
            },
            method="GET",
        )

        try:
            with urllib.request.urlopen(req, timeout=10) as resp:
                raw = resp.read().decode("utf-8")
        except Exception:
            return []

        try:
            rows = json.loads(raw)
        except Exception:
            return []

        return rows if isinstance(rows, list) else []

    rows = _request_rows()
    items: List[Tuple[str, WD_COLOR_INDEX]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        row_client = (row.get("Client") or row.get("client") or "").strip()
        if client.lower() not in row_client.lower():
            continue
        word = (row.get("Word") or "").strip()
        color = _color_from_supabase(row.get("HighlightColor") or "")
        if word and color:
            items.append((word, color))

    if "by_client" not in _HIGHLIGHT_CACHE or not isinstance(_HIGHLIGHT_CACHE.get("by_client"), dict):
        _HIGHLIGHT_CACHE["by_client"] = {}
    _HIGHLIGHT_CACHE["by_client"][client] = {"ts": now, "items": items}
    return items


def _find_highlight_spans(text: str, words: List[Tuple[str, WD_COLOR_INDEX]]) -> List[HighlightSpan]:
    if not text or not words:
        return []

    spans: List[HighlightSpan] = []
    for word, color in words:
        escaped = re.escape(word)
        # If the word is alphanumeric-ish, treat it as a whole word; otherwise match as-is.
        if re.fullmatch(r"[A-Za-z0-9']+", word):
            pattern = re.compile(rf"\b{escaped}\b", re.IGNORECASE)
        else:
            pattern = re.compile(escaped, re.IGNORECASE)

        for m in pattern.finditer(text):
            spans.append(HighlightSpan(start=m.start(), end=m.end(), color=color))

    spans.sort(key=lambda s: (s.start, s.end))
    return spans


def _add_formatted_runs(paragraph, text: str, highlight_spans: List[HighlightSpan]) -> None:
    if text is None:
        text = ""

    diary_spans = [m.span() for m in _DIARY_PATTERN.finditer(text)]
    budget_spans = [m.span() for m in _BUDGET_PATTERN.finditer(text)]

    boundaries = {0, len(text)}
    for s, e in diary_spans:
        boundaries.add(s)
        boundaries.add(e)
    for s, e in budget_spans:
        boundaries.add(s)
        boundaries.add(e)
    for sp in highlight_spans:
        boundaries.add(sp.start)
        boundaries.add(sp.end)

    points = sorted(boundaries)

    def _in_diary(i: int) -> bool:
        for s, e in diary_spans:
            if s <= i < e:
                return True
        return False

    def _in_budget(i: int) -> bool:
        for s, e in budget_spans:
            if s <= i < e:
                return True
        return False

    def _highlight_at(i: int) -> Optional[WD_COLOR_INDEX]:
        for sp in highlight_spans:
            if sp.start <= i < sp.end:
                return sp.color
        return None

    for a, b in zip(points, points[1:]):
        if a == b:
            continue
        seg = text[a:b]
        r = paragraph.add_run(seg)
        _set_font(r, name="Times New Roman", size_pt=9, bold=False)

        if _in_diary(a) or _in_budget(a):
            r.italic = True

        color = _highlight_at(a)
        if color is not None:
            r.font.highlight_color = color


def generate_published_letter_docx_bytes(
    *,
    letter_heading: Optional[str],
    letter_subheading: Optional[str],
    date: Optional[str],
    body: Any,
    author_name: Optional[str],
) -> bytes:
    heading_line = _normalize_location_line(letter_heading or "")
    subheading_line = (letter_subheading or "").strip()
    date_text = (date or "").strip()
    author = (author_name or "").strip()

    paragraphs = _split_paragraphs(body)

    # Normalize Diary/The Diary and Budget/The Budget before rendering.
    paragraphs = [
        normalize_budget_casing(normalize_diary_casing(strip_diary_markup(strip_budget_markup(p))))
        for p in paragraphs
    ]

    spellcheck_words = _fetch_spellcheck_words_for_client("Budget")

    doc = Document()

    # Heading
    if heading_line:
        p = doc.add_paragraph()
        p.style = "No Spacing"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(heading_line)
        _set_font(r, name="Helvetica", size_pt=10, bold=True)

    # Subheading
    if subheading_line:
        p = doc.add_paragraph()
        p.style = "No Spacing"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(subheading_line)
        _set_font(r, name="Times New Roman", size_pt=9, bold=True)

    # Body
    if date_text and paragraphs:
        # Date + EM dash + opening sentence in same paragraph.
        first = f"{date_text}—{paragraphs[0].lstrip()}"
        paragraphs = [first] + paragraphs[1:]

    for idx, para_text in enumerate(paragraphs or [""]):
        p = doc.add_paragraph()
        p.style = "No Spacing"
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.first_line_indent = Inches(0.15)

        spans = _find_highlight_spans(para_text, spellcheck_words)
        _add_formatted_runs(p, para_text, spans)

    # Author name bottom-right
    if author:
        p = doc.add_paragraph()
        p.style = "No Spacing"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(author)
        _set_font(r, name="Times New Roman", size_pt=9, bold=False)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
